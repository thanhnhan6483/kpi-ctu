import docx
import sys
import os
import re

def docx_to_markdown(docx_path, md_path):
    doc = docx.Document(docx_path)
    md_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue
        
        style = para.style.name if para.style else ""
        
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading", "").strip())
            except ValueError:
                level = 1
            md_lines.append(f"{'#' * level} {text}")
        else:
            formatted = ""
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold and run.italic:
                    t = f"***{t}***"
                elif run.bold:
                    t = f"**{t}**"
                elif run.italic:
                    t = f"*{t}*"
                formatted += t
            if formatted:
                md_lines.append(formatted)
            else:
                md_lines.append(text)
    
    for table in doc.tables:
        md_lines.append("")
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip())
        
        md_lines.append("| " + " | ".join(headers) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")
    
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    
    print(f"Converted: {docx_path} -> {md_path}", flush=True)

if __name__ == "__main__":
    kpi_dir = sys.argv[1] if len(sys.argv) > 1 else "."
    for f in os.listdir(kpi_dir):
        if f.endswith(".docx") and not f.startswith(".~lock"):
            docx_path = os.path.join(kpi_dir, f)
            md_path = os.path.join(kpi_dir, f.replace(".docx", ".md"))
            if not os.path.exists(md_path):
                docx_to_markdown(docx_path, md_path)
